"""
Word document converter (DOCX format)
"""
from docx import Document
from typing import Dict, Any, List
from .base_converter import BaseConverter
import datetime


class WordConverter(BaseConverter):
    """Convert Word documents to structured data"""
    
    def __init__(self, file_path: str):
        super().__init__(file_path)
        self.doc = None
        
    def extract(self) -> Dict[str, Any]:
        """Extract text and structure from Word document"""
        try:
            self.doc = Document(self.file_path)
            
            paragraphs = []
            full_text = []
            
            for para in self.doc.paragraphs:
                if para.text.strip():
                    para_data = {
                        "text": para.text,
                        "style": para.style.name,
                        "word_count": len(para.text.split())
                    }
                    paragraphs.append(para_data)
                    full_text.append(para.text)
            
            # Extract tables
            tables = self._extract_tables()
            
            result = {
                "document_type": "Word Document",
                "total_paragraphs": len(paragraphs),
                "total_tables": len(tables),
                "full_text": "\n\n".join(full_text),
                "paragraphs": paragraphs,
                "tables": tables,
                "metadata": self.get_metadata()
            }
            
            return result
            
        except Exception as e:
            return {
                "error": f"Failed to extract Word document: {str(e)}",
                "document_type": "Word Document"
            }
    
    def _extract_tables(self) -> List[Dict[str, Any]]:
        """Extract tables from Word document"""
        tables_data = []
        
        try:
            for table_idx, table in enumerate(self.doc.tables):
                rows_data = []
                
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    rows_data.append(row_data)
                
                table_data = {
                    "table_index": table_idx + 1,
                    "rows": len(rows_data),
                    "columns": len(rows_data[0]) if rows_data else 0,
                    "data": rows_data
                }
                tables_data.append(table_data)
        except Exception as e:
            pass
            
        return tables_data
    
    def get_metadata(self) -> Dict[str, Any]:
        """Get Word document metadata"""
        if not self.doc:
            self.doc = Document(self.file_path)
        
        core_props = self.doc.core_properties
        
        return {
            "title": core_props.title or "N/A",
            "author": core_props.author or "N/A",
            "subject": core_props.subject or "N/A",
            "keywords": core_props.keywords or "N/A",
            "created": str(core_props.created) if core_props.created else "N/A",
            "modified": str(core_props.modified) if core_props.modified else "N/A",
            "last_modified_by": core_props.last_modified_by or "N/A"
        }
