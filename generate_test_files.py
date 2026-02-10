"""
Generate test files for AI Data Conversion Tool
Creates sample PDF, Word, Excel, TXT, and CSV files
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
import os

# Create test_data directory
os.makedirs('test_data', exist_ok=True)

print("🔧 Generating test files for AI Data Conversion Tool...")
print("=" * 60)

# ============================================================================
# 1. Generate PDF - Sample Invoice
# ============================================================================
print("\n📄 Creating PDF file (sample_invoice.pdf)...")

pdf_file = "test_data/sample_invoice.pdf"
doc = SimpleDocTemplate(pdf_file, pagesize=letter)
elements = []
styles = getSampleStyleSheet()

# Title
title_style = ParagraphStyle(
    'CustomTitle',
    parent=styles['Heading1'],
    fontSize=24,
    textColor=colors.HexColor('#1a5490'),
    spaceAfter=30,
    alignment=1  # Center
)
elements.append(Paragraph("INVOICE", title_style))
elements.append(Spacer(1, 0.3*inch))

# Invoice details
invoice_details = [
    ["Invoice Number:", "#INV-2026-001"],
    ["Date:", "February 10, 2026"],
    ["Customer:", "Acme Corporation"],
    ["Email:", "billing@acmecorp.com"],
]

detail_table = Table(invoice_details, colWidths=[2*inch, 3*inch])
detail_table.setStyle(TableStyle([
    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
    ('FONTSIZE', (0, 0), (-1, -1), 11),
    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
]))
elements.append(detail_table)
elements.append(Spacer(1, 0.4*inch))

# Items table
elements.append(Paragraph("Items and Services", styles['Heading2']))
elements.append(Spacer(1, 0.2*inch))

items_data = [
    ["Item", "Description", "Quantity", "Unit Price", "Total"],
    ["Cloud Storage", "Premium 1TB Storage", "1", "$50.00", "$50.00"],
    ["API Access", "Unlimited API Calls", "1", "$100.00", "$100.00"],
    ["Support Package", "24/7 Premium Support", "1", "$75.00", "$75.00"],
    ["Data Processing", "ML Model Training", "5", "$30.00", "$150.00"],
]

items_table = Table(items_data, colWidths=[1.5*inch, 2*inch, 1*inch, 1*inch, 1*inch])
items_table.setStyle(TableStyle([
    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a5490')),
    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
    ('FONTSIZE', (0, 0), (-1, 0), 11),
    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
    ('GRID', (0, 0), (-1, -1), 1, colors.black),
]))
elements.append(items_table)
elements.append(Spacer(1, 0.3*inch))

# Totals
totals_data = [
    ["Subtotal:", "$375.00"],
    ["Tax (10%):", "$37.50"],
    ["Total:", "$412.50"],
]

totals_table = Table(totals_data, colWidths=[5*inch, 1.5*inch])
totals_table.setStyle(TableStyle([
    ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
    ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
    ('FONTSIZE', (0, -1), (-1, -1), 14),
    ('LINEABOVE', (0, -1), (-1, -1), 2, colors.black),
    ('TOPPADDING', (0, -1), (-1, -1), 10),
]))
elements.append(totals_table)

# Footer
elements.append(Spacer(1, 0.5*inch))
footer_text = "Thank you for your business! Payment is due within 30 days. For questions, contact support@aiconverter.com"
elements.append(Paragraph(footer_text, styles['Normal']))

doc.build(elements)
print("✅ PDF created successfully!")

# ============================================================================
# 2. Generate Word Document - Project Report
# ============================================================================
print("\n📝 Creating Word document (project_report.docx)...")

word_file = "test_data/project_report.docx"
doc = Document()

# Title
title = doc.add_heading('AI Data Conversion Tool - Project Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadata
doc.add_paragraph('Date: February 10, 2026')
doc.add_paragraph('Author: Research Team')
doc.add_paragraph('Department: AI & Machine Learning')
doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', 1)
summary_text = """
This report presents the findings of our research on automated data conversion tools 
for artificial intelligence training. The AI Data Conversion Tool successfully converts 
multiple file formats (PDF, Word, Excel, TXT) into structured datasets with 95% accuracy. 
Our testing shows an 80% reduction in data preparation time compared to manual methods.
"""
doc.add_paragraph(summary_text)

# Methodology
doc.add_heading('Methodology', 1)
methodology_text = """
We conducted a comprehensive evaluation using the following approach:
"""
doc.add_paragraph(methodology_text)

# Add bullet points
doc.add_paragraph('Tested 1,000+ documents across 5 file formats', style='List Bullet')
doc.add_paragraph('Measured extraction accuracy using ground truth datasets', style='List Bullet')
doc.add_paragraph('Compared processing time against manual conversion', style='List Bullet')
doc.add_paragraph('Evaluated output quality for ML model training', style='List Bullet')

# Results Table
doc.add_heading('Results', 1)
doc.add_paragraph('Performance metrics across different file types:')

table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = 'File Type'
header_cells[1].text = 'Accuracy'
header_cells[2].text = 'Avg Time (sec)'
header_cells[3].text = 'Success Rate'

# Data rows
data = [
    ['PDF', '94%', '2.3', '98%'],
    ['Word (DOCX)', '96%', '1.8', '99%'],
    ['Excel (XLSX)', '97%', '1.5', '99%'],
    ['Text/CSV', '98%', '0.8', '100%'],
    ['Average', '96.25%', '1.6', '99%']
]

for i, row_data in enumerate(data, start=1):
    cells = table.rows[i].cells
    for j, value in enumerate(row_data):
        cells[j].text = value

# Conclusion
doc.add_heading('Conclusion', 1)
conclusion_text = """
The AI Data Conversion Tool demonstrates exceptional performance across all tested 
file formats. With an average accuracy of 96.25% and processing time of 1.6 seconds 
per file, it significantly outperforms manual conversion methods. The tool is ready 
for production deployment and can handle enterprise-scale data preparation workflows.

Key achievements:
- 80% reduction in data preparation time
- 96%+ extraction accuracy across all formats
- Support for batch processing of multiple files
- Production-ready with comprehensive error handling
"""
doc.add_paragraph(conclusion_text)

# Save document
doc.save(word_file)
print("✅ Word document created successfully!")

# ============================================================================
# 3. Generate Excel Spreadsheet - Sales Data
# ============================================================================
print("\n📊 Creating Excel file (sales_data.xlsx)...")

excel_file = "test_data/sales_data.xlsx"

# Create main sales data
sales_data = {
    'Order ID': ['ORD-001', 'ORD-002', 'ORD-003', 'ORD-004', 'ORD-005', 
                 'ORD-006', 'ORD-007', 'ORD-008', 'ORD-009', 'ORD-010'],
    'Customer Name': ['John Smith', 'Sarah Johnson', 'Mike Williams', 'Emily Brown', 'David Lee',
                     'Lisa Anderson', 'James Wilson', 'Maria Garcia', 'Robert Taylor', 'Jennifer White'],
    'Product': ['Laptop', 'Mouse', 'Keyboard', 'Monitor', 'Headphones',
               'Webcam', 'USB Hub', 'External SSD', 'Laptop Stand', 'Cable Organizer'],
    'Quantity': [2, 5, 3, 1, 4, 2, 6, 1, 3, 10],
    'Unit Price': [899.99, 29.99, 79.99, 449.99, 89.99, 
                   149.99, 34.99, 199.99, 49.99, 12.99],
    'Total': [1799.98, 149.95, 239.97, 449.99, 359.96,
             299.98, 209.94, 199.99, 149.97, 129.90],
    'Date': ['2026-02-01', '2026-02-02', '2026-02-03', '2026-02-04', '2026-02-05',
            '2026-02-06', '2026-02-07', '2026-02-08', '2026-02-09', '2026-02-10'],
    'Status': ['Completed', 'Completed', 'Pending', 'Completed', 'Shipped',
              'Completed', 'Processing', 'Completed', 'Shipped', 'Completed']
}

df_sales = pd.DataFrame(sales_data)

# Create summary statistics sheet
summary_data = {
    'Metric': ['Total Orders', 'Total Revenue', 'Average Order Value', 'Completed Orders', 'Pending Orders'],
    'Value': [10, '$3989.63', '$398.96', 6, 1]
}
df_summary = pd.DataFrame(summary_data)

# Create product categories sheet
categories_data = {
    'Category': ['Electronics', 'Accessories', 'Peripherals', 'Storage', 'Organization'],
    'Items': [4, 3, 2, 1, 1],
    'Revenue': ['$2499.91', '$699.89', '$389.92', '$199.99', '$129.90']
}
df_categories = pd.DataFrame(categories_data)

# Write to Excel with multiple sheets
with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
    df_sales.to_excel(writer, sheet_name='Sales Data', index=False)
    df_summary.to_excel(writer, sheet_name='Summary', index=False)
    df_categories.to_excel(writer, sheet_name='Categories', index=False)

print("✅ Excel file created successfully!")

# ============================================================================
# 4. Generate Text File - Meeting Notes
# ============================================================================
print("\n📋 Creating text file (meeting_notes.txt)...")

txt_file = "test_data/meeting_notes.txt"

txt_content = """AI DATA CONVERSION TOOL - TEAM MEETING NOTES
=====================================================

Date: February 10, 2026
Time: 10:00 AM - 11:30 AM
Location: Conference Room A
Attendees: Sarah (PM), John (Dev), Emily (QA), Mike (Design)

AGENDA ITEMS
------------

1. Project Status Update
   - Development completed on schedule
   - All core features implemented
   - Testing in progress

2. Feature Highlights
   - Multi-format support (PDF, Word, Excel, TXT, CSV)
   - Batch processing capability
   - Real-time preview and statistics
   - AI-optimized output formats

3. Testing Results
   - 96% accuracy across all formats
   - Average processing time: 1.6 seconds
   - Batch processing: 50 files in under 2 minutes
   - Zero critical bugs found

4. User Feedback
   - Interface is intuitive and user-friendly
   - Preview feature highly appreciated
   - Batch processing saves significant time
   - Documentation is comprehensive

5. Next Steps
   - Complete final QA testing
   - Prepare demo for professor
   - Create presentation materials
   - Schedule deployment review

ACTION ITEMS
------------
- Sarah: Finalize presentation slides by Feb 11
- John: Fix minor UI alignment issues by Feb 10
- Emily: Complete testing report by Feb 11
- Mike: Update user guide with screenshots by Feb 11

TECHNICAL NOTES
---------------
- Backend: Python 3.8+ with specialized libraries
- Frontend: Gradio framework for rapid UI development
- PDF Processing: PyMuPDF + pdfplumber for accuracy
- Data Format: Support for JSON, CSV, XML outputs

DISCUSSION POINTS
-----------------
The team agreed that the AI Training Format output is a unique 
feature that differentiates our tool from competitors. This format 
includes metadata, features, and structured data specifically 
optimized for machine learning workflows.

Performance testing showed impressive results with large files. 
A 50MB PDF was processed in under 8 seconds, well within acceptable 
limits. Batch processing of 100 files completed in approximately 
3 minutes.

RISKS AND MITIGATION
---------------------
- Risk: Some PDFs with complex layouts may have extraction issues
  Mitigation: Using dual-library approach (PyMuPDF + pdfplumber)

- Risk: Very large Excel files may cause memory issues
  Mitigation: Implemented streaming and file size limits

NEXT MEETING
------------
Date: February 15, 2026
Time: 2:00 PM
Topic: Post-submission review and future enhancements

NOTES
-----
Overall, the team is confident in the project's quality and 
readiness for submission. All core objectives have been met, 
and the tool is production-ready with comprehensive documentation.

Meeting adjourned at 11:30 AM.

---
End of Meeting Notes
"""

with open(txt_file, 'w', encoding='utf-8') as f:
    f.write(txt_content)

print("✅ Text file created successfully!")

# ============================================================================
# 5. Generate CSV File - Customer Database
# ============================================================================
print("\n📈 Creating CSV file (customer_database.csv)...")

csv_file = "test_data/customer_database.csv"

csv_data = {
    'Customer_ID': ['CUST-001', 'CUST-002', 'CUST-003', 'CUST-004', 'CUST-005',
                   'CUST-006', 'CUST-007', 'CUST-008', 'CUST-009', 'CUST-010'],
    'First_Name': ['Alice', 'Bob', 'Charlie', 'Diana', 'Edward',
                   'Fiona', 'George', 'Hannah', 'Ian', 'Julia'],
    'Last_Name': ['Johnson', 'Smith', 'Williams', 'Brown', 'Jones',
                  'Davis', 'Miller', 'Wilson', 'Moore', 'Taylor'],
    'Email': ['alice.johnson@email.com', 'bob.smith@email.com', 'charlie.williams@email.com',
             'diana.brown@email.com', 'edward.jones@email.com', 'fiona.davis@email.com',
             'george.miller@email.com', 'hannah.wilson@email.com', 'ian.moore@email.com',
             'julia.taylor@email.com'],
    'Phone': ['555-0101', '555-0102', '555-0103', '555-0104', '555-0105',
             '555-0106', '555-0107', '555-0108', '555-0109', '555-0110'],
    'Company': ['Tech Corp', 'Data Systems', 'Cloud Solutions', 'AI Innovations', 'Smart Analytics',
               'Future Tech', 'Digital Experts', 'Code Masters', 'Info Systems', 'Tech Pioneers'],
    'Industry': ['Technology', 'IT Services', 'Cloud Computing', 'AI/ML', 'Data Analytics',
                'Software', 'Consulting', 'Development', 'Systems', 'Research'],
    'Registration_Date': ['2025-01-15', '2025-02-20', '2025-03-10', '2025-04-05', '2025-05-12',
                         '2025-06-18', '2025-07-22', '2025-08-30', '2025-09-14', '2025-10-25'],
    'Subscription_Type': ['Premium', 'Basic', 'Premium', 'Enterprise', 'Basic',
                         'Premium', 'Basic', 'Enterprise', 'Premium', 'Basic'],
    'Status': ['Active', 'Active', 'Active', 'Active', 'Inactive',
              'Active', 'Active', 'Active', 'Pending', 'Active']
}

df_csv = pd.DataFrame(csv_data)
df_csv.to_csv(csv_file, index=False, encoding='utf-8')

print("✅ CSV file created successfully!")

# ============================================================================
# Summary
# ============================================================================
print("\n" + "=" * 60)
print("🎉 TEST FILE GENERATION COMPLETE!")
print("=" * 60)
print("\n📁 Created files in 'test_data/' directory:")
print("   1. ✅ sample_invoice.pdf - Invoice with tables and text")
print("   2. ✅ project_report.docx - Document with headings and tables")
print("   3. ✅ sales_data.xlsx - Multi-sheet Excel workbook")
print("   4. ✅ meeting_notes.txt - Plain text meeting notes")
print("   5. ✅ customer_database.csv - CSV customer database")
print("\n💡 You can now upload these files to test the application!")
print("   Run: python app.py")
print("   Then upload files from the test_data/ folder")
print("\n" + "=" * 60)
